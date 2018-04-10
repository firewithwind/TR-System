#-*-coding:utf-8-*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.oxml.ns import qn

data = [
    [u'申请人', u'刘连兴', u'目的地', u'北京'],
    [u'所属项目', u'重点项目', u'申请时间', u'2017/01/02'],
    [u'项目描述', u'11111', u'11111', u'11111'],
    [u'项目描述', u'11111', u'11111', u'11111'],
    [u'项目描述', u'11111', u'11111', u'11111'],
    [u'项目描述', u'11111', u'11111', u'11111']
]

doc = Document()

doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal'].font.size = Pt(14)

header = doc.add_paragraph()
header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run(u'网络与信息安全技术研究中心出差任务审批单')
run.bold = True

table = doc.add_table(rows=6, cols=4, style = 'Table Grid')
table.style.font.size = Pt(14)
# table.style.next_paragraph_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

table.cell(2, 1).merge(table.cell(2, 3))
table.cell(3, 1).merge(table.cell(3, 3))
table.cell(4, 1).merge(table.cell(4, 3))
table.cell(5, 1).merge(table.cell(5, 3))

# row = table.rows[0]
# tr = row._tr
# trHeight = tr.xpath('./w:trPr/w:trHeight')
# trHeight.set(qn('w:val'), "2092")
# trHeight.set(qn('w:hRule'), "atLeast")
print trHeight

for i in range(0, 6):
    for j in range(0, 4):
        table.cell(i, j).text = data[i][j]
        if j % 2 == 0:
            table.cell(i, j).width = Cm(2.74)
        elif i < 2:
            table.cell(i, j).width = Cm(4)
        else:
            table.cell(i, j).width = Cm(9.75)
print table.style

doc.save('test.docx')
